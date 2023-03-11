"""
 Copyright (C) 2022 Fern Lane, Webinar-hacker
 Licensed under the GNU Affero General Public License, Version 3.0 (the "License");
 you may not use this file except in compliance with the License.
 You may obtain a copy of the License at
       https://www.gnu.org/licenses/agpl-3.0.en.html
 Unless required by applicable law or agreed to in writing, software
 distributed under the License is distributed on an "AS IS" BASIS,
 WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
 See the License for the specific language governing permissions and
 limitations under the License.
 IN NO EVENT SHALL THE AUTHOR BE LIABLE FOR ANY CLAIM, DAMAGES OR
 OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
 ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR
 OTHER DEALINGS IN THE SOFTWARE.
"""

import logging
import os
import threading
import time

from docx import Document
from docx.shared import Inches, RGBColor, Pt

from BrowserHandler import SCREENSHOT_EXTENSION

WAVE_FILE_SIZE_MIN_BYTES = 100


class LectureBuilder:
    def __init__(self, settings, elements_set_enabled_signal, progress_bar_set_value_signal,
                 progress_bar_set_maximum_signal, lecture_copy_signal, label_device_signal,
                 label_time_left_signal):
        self.settings = settings
        self.elements_set_enabled_signal = elements_set_enabled_signal
        self.progress_bar_set_value_signal = progress_bar_set_value_signal
        self.progress_bar_set_maximum_signal = progress_bar_set_maximum_signal
        self.lecture_copy_signal = lecture_copy_signal
        self.label_device_signal = label_device_signal
        self.label_time_left_signal = label_time_left_signal

        self.audio_files = []
        self.screenshots = []
        self.audio_bytes_total = 0
        self.lecture_name = ''
        self.model = None

    def start_building_lecture(self, lecture_directory: str, lecture_name: str):
        """
        Starts building lecture
        :param lecture_name: example DD_MM_YYYY__HH_MM_SS
        :param lecture_directory: example recordings/DD_MM_YYYY__HH_MM_SS
        :return:
        """
        logging.info('Building lecture ' + lecture_name)
        self.lecture_name = lecture_name
        self.audio_files = []
        self.screenshots = []
        self.audio_bytes_total = 0

        # Find audio files
        for audio_or_screenshot_dir in os.listdir(lecture_directory):
            audio_or_screenshot_dir = os.path.join(lecture_directory, audio_or_screenshot_dir)
            if os.path.isdir(audio_or_screenshot_dir) \
                    and str(self.settings['audio_directory_name']) in audio_or_screenshot_dir:
                for file_ in os.listdir(audio_or_screenshot_dir):
                    file_ = os.path.join(audio_or_screenshot_dir, file_)
                    if not os.path.isdir(file_) and '.wav' in str(file_):
                        time_diff = ''.join(os.path.basename(file_).strip().split('.')[: -1])
                        time_diff_int = -1
                        try:
                            time_diff_int = int(time_diff)
                        except:
                            pass
                        if time_diff_int >= 0:
                            logging.info('Found audio file: ' + str(file_) + ' with time: ' + str(time_diff_int))

                            # Check file size
                            file_size = os.path.getsize(str(file_))
                            if file_size < WAVE_FILE_SIZE_MIN_BYTES:
                                logging.warning('Size of file ' + str(file_) + ' too small! Ignoring it')
                            else:
                                self.audio_files.append([time_diff_int, str(file_), file_size])
                                self.audio_bytes_total += file_size

        # Find screenshots
        screenshots_dir = os.path.join(lecture_directory, str(self.settings['screenshots_directory_name']))
        if os.path.exists(screenshots_dir):
            for file in os.listdir(screenshots_dir):
                dir_or_file = os.path.join(screenshots_dir, file)
                if not os.path.isdir(dir_or_file) and SCREENSHOT_EXTENSION in str(dir_or_file):
                    time_diff = ''.join(os.path.basename(dir_or_file).strip().split('.')[: -1])
                    time_diff_int = -1
                    try:
                        time_diff_int = int(time_diff)
                    except:
                        pass
                    if time_diff_int >= 0:
                        logging.info('Found screenshot: ' + str(dir_or_file) + ' with time: ' + str(time_diff_int))
                        self.screenshots.append([time_diff_int, str(dir_or_file)])

        # Sort audio files and screenshots
        if len(self.audio_files) > 0:
            self.audio_files.sort(key=lambda x: x[0])
        if len(self.screenshots) > 0:
            self.screenshots.sort(key=lambda x: x[0], reverse=True)

        # Check for audio file
        if len(self.audio_files) > 0:
            # Start thread
            thread = threading.Thread(target=self.lecture_builder_thread)
            thread.start()
            logging.info('Lecture builder thread: ' + thread.name)

        # No audio file
        else:
            logging.warning('No audio file!')
            # Enable gui elements
            self.elements_set_enabled_signal.emit(True)

    def lecture_builder_thread(self):
        """
        Transcribes audio and build lecture
        :return:
        """
        try:
            # Load packages
            logging.info('Importing packages...')
            import whisper_timestamped as whisper

            # Load model
            if self.model is None:
                # Select cpu or gpu
                import torch
                device = 'cuda' if torch.cuda.is_available() else 'cpu'
                logging.info('Device: ' + device)
                self.label_device_signal.emit('Device: ' + device)

                # Load model
                model_dir = os.getcwd()
                logging.info('Loading model into: ' + model_dir)
                self.model = whisper.load_model(str(self.settings['whisper_model_name']), device=device,
                                                download_root=model_dir)

            # Result lists
            words = []
            timestamps_end = []
            confidences_percents = []

            # Transcribe audio
            logging.info('Starting transcription... Please wait')
            seconds_per_byte_filtered = 0
            self.progress_bar_set_maximum_signal.emit(len(self.audio_files))
            self.label_time_left_signal.emit('Time left: 00:00:00')
            for audio_file_n in range(len(self.audio_files)):
                # Set progress
                self.progress_bar_set_value_signal.emit(audio_file_n + 1)

                transcription = None
                audio_file_ = self.audio_files[audio_file_n]
                try:
                    # Record start time
                    transcription_time_started = time.time()

                    # Load audio file
                    audio = whisper.load_audio(audio_file_[1])
                    audio = whisper.pad_or_trim(audio)

                    # Transcribe audio
                    transcription = whisper.transcribe(self.model, audio,
                                                       language=str(self.settings['whisper_model_language']))

                    # Calculate seconds per byte
                    seconds_per_byte = (time.time() - transcription_time_started) / audio_file_[2]
                    if seconds_per_byte_filtered == 0:
                        seconds_per_byte_filtered = seconds_per_byte
                    else:
                        filter_factor = float(self.settings['lecture_build_time_filter_factor'])
                        seconds_per_byte_filtered = seconds_per_byte_filtered * filter_factor \
                                                    + seconds_per_byte * (1. - filter_factor)
                    logging.info('Microseconds per byte: ' + str(int(seconds_per_byte * 1000 * 1000)) + ', avg: '
                                 + str(int(seconds_per_byte_filtered * 1000 * 1000)))

                    # Subtract processed bytes
                    self.audio_bytes_total -= audio_file_[2]

                    # Calculate and show time left
                    seconds_left = self.audio_bytes_total * seconds_per_byte_filtered
                    logging.info('Time left: ~' + str(int(seconds_left)) + 's')
                    time_left_seconds = int(seconds_left % 60)
                    time_left_minutes = int((seconds_left / 60) % 60)
                    time_left_hours = int(seconds_left / (60 * 60))
                    self.label_time_left_signal.emit('Time left: ' + '{:02d}'.format(time_left_hours)
                                                     + ':' + '{:02d}'.format(time_left_minutes)
                                                     + ':' + '{:02d}'.format(time_left_seconds))
                # Error
                except Exception as e:
                    logging.warning(e)

                # Parse result
                if transcription is not None and transcription['segments'] is not None \
                        and len(transcription['segments']) > 0:
                    for segment in transcription['segments']:
                        if segment is not None and segment['words'] is not None and len(segment['words']) > 0:
                            for segment_word in segment['words']:
                                if segment_word is not None:
                                    if segment_word['text'] is not None and segment_word['end'] is not None \
                                            and segment_word['confidence'] is not None:
                                        text_ = str(segment_word['text']).strip()
                                        if len(text_) > 0:
                                            # Finally, append data
                                            words.append(text_)
                                            timestamps_end.append(int(1000. * float(segment_word['end']))
                                                                  + audio_file_[0])
                                            confidences_percents.append(int(100. * float(segment_word['confidence'])))

            # Log length of words
            logging.info('Transcription result words: ' + str(len(words)))

            # Check size of words list
            if len(words) > 0:
                # Build docx
                self.write_to_docx(words, timestamps_end, confidences_percents)

                # Done
                self.lecture_copy_signal.emit(os.path.join(self.settings['lectures_directory_name'],
                                                           self.lecture_name + '.docx'))

            # No words
            else:
                logging.warning('No words to write!')

        # Error building lecture
        except Exception as e:
            logging.error(e, exc_info=True)

        # Reset progress
        self.progress_bar_set_maximum_signal.emit(100)
        self.progress_bar_set_value_signal.emit(0)
        self.label_time_left_signal.emit('Time left: 00:00:00')

        # Enable gui elements
        self.elements_set_enabled_signal.emit(True)

    def write_to_docx(self, words: list, timestamps_end: list, confidences_percents: list):
        """
        Finally writes words and screenshots to docx document
        :param words:
        :param timestamps_end:
        :param confidences_percents:
        :return:
        """
        # Create docx document
        logging.info('Writing to docx document...')
        document = Document()
        document.add_heading(self.lecture_name, 0)

        # First screenshot
        current_screenshot = None
        if len(self.screenshots) > 0:
            current_screenshot = self.screenshots.pop()

        # Reset progress
        self.progress_bar_set_maximum_signal.emit(len(words))
        self.progress_bar_set_value_signal.emit(0)

        # Create initial paragraph
        paragraph = document.add_paragraph('')

        # Get initial timestamp
        timestamp_last = timestamps_end[0]

        # List all words
        for word_n in range(len(words)):
            # Set progress
            self.progress_bar_set_value_signal.emit(word_n + 1)

            # Extract data
            word = str(words[word_n])
            timestamp_end = timestamps_end[word_n]
            confidence_percents = confidences_percents[word_n]

            # New paragraph
            if timestamp_end - timestamp_last >= int(self.settings['paragraph_audio_distance_min_milliseconds']):
                paragraph = document.add_paragraph('')
            timestamp_last = timestamp_end

            # Add screenshots
            while current_screenshot is not None and timestamp_end >= current_screenshot[0]:
                # New paragraph
                document.add_paragraph('')

                # Append screenshot
                document.add_picture(os.path.normpath(str(current_screenshot[1])), width=Inches(
                    float(self.settings['lecture_picture_width_inches'])))

                # Get next screenshot
                if len(self.screenshots) > 0:
                    current_screenshot = self.screenshots.pop()
                else:
                    current_screenshot = None

                # New paragraph
                paragraph = document.add_paragraph('')

            # Append word
            run_ = paragraph.add_run(word + ' ')

            # Set font size
            run_.font.size = Pt(int(self.settings['lecture_font_size_pt']))

            # Show low probability words
            if confidence_percents <= int(self.settings['word_low_confidence_threshold_percents']):
                text_colors = self.settings['lecture_low_confidence_text_color']
                run_.font.color.rgb = RGBColor(int(text_colors[0]),
                                               int(text_colors[1]),
                                               int(text_colors[2]))
            else:
                text_colors = self.settings['lecture_default_text_color']
                run_.font.color.rgb = RGBColor(int(text_colors[0]),
                                               int(text_colors[1]),
                                               int(text_colors[2]))

        # Add all remaining screenshots
        while len(self.screenshots) > 0:
            # Pop screenshot
            current_screenshot = self.screenshots.pop()

            # New paragraph
            document.add_paragraph('')

            # Append screenshot
            document.add_picture(os.path.normpath(str(current_screenshot[1])), width=Inches(
                float(self.settings['lecture_picture_width_inches'])))

        # Create lectures directory
        lectures_dir = str(self.settings['lectures_directory_name'])
        if not os.path.exists(lectures_dir):
            os.makedirs(lectures_dir)

        # Save lecture
        lecture_file = os.path.join(lectures_dir, self.lecture_name + '.docx')
        logging.info('Saving lecture as: ' + lecture_file)
        document.save(lecture_file)
